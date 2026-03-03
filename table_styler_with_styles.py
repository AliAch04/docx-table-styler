"""
DOCX Table Style Changer - Using your document's actual styles
"""

import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor

def get_all_tables_recursive(element, tables_list=None):
    """Recursively find ALL tables including nested ones"""
    if tables_list is None:
        tables_list = []
    
    if hasattr(element, 'tables'):
        for table in element.tables:
            if table not in tables_list:
                tables_list.append(table)
                # Recursively search inside this table's cells
                for row in table.rows:
                    for cell in row.cells:
                        get_all_tables_recursive(cell, tables_list)
    return tables_list

def apply_best_style(table, style_name, available_styles):
    """
    Apply the closest matching style available in the document
    """
    # If the requested style is available, use it
    if style_name in available_styles:
        try:
            table.style = style_name
            return True, f"Applied '{style_name}'"
        except:
            pass
    
    # Try to find a similar style
    style_lower = style_name.lower()
    for avail in available_styles:
        if style_lower in avail.lower() or avail.lower() in style_lower:
            try:
                table.style = avail
                return True, f"Applied similar style: '{avail}'"
            except:
                continue
    
    # If no match found, use the first available table style
    if available_styles:
        try:
            table.style = available_styles[0]
            return True, f"Applied default style: '{available_styles[0]}'"
        except:
            pass
    
    return False, "No applicable style found"

def apply_manual_formatting(table):
    """Apply nice formatting manually if styles fail"""
    try:
        # Set table properties
        table.autofit = False
        table.allow_autofit = True
        
        # Format header row
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                # Add shading to header
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                   r'w:val="clear" w:color="auto" w:fill="D9E1F2"/>')
                tcPr.append(shading)
                
                # Make header text bold
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
        
        # Add borders to all cells
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add simple borders
                borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                   r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                   r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                   r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                   r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
                                   r'</w:tcBorders>')
                tcPr.append(borders)
        
        return True, "Manual formatting applied"
    except Exception as e:
        return False, f"Manual formatting failed: {e}"

def process_document(input_path, output_path, style_choice="auto"):
    """
    Process the document with appropriate styling
    style_choice can be: "auto", "normal", "plain", or specific style name
    """
    print(f"📂 Loading: {input_path}")
    doc = Document(input_path)
    
    # Get available table styles
    available_styles = [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE]
    print(f"📋 Available table styles in this document: {available_styles}")
    
    # Determine which style to use
    style_map = {
        "auto": available_styles[0] if available_styles else None,
        "normal": "Normal Table" if "Normal Table" in available_styles else None,
        "plain": "Plain Table 3" if "Plain Table 3" in available_styles else None,
    }
    
    target_style = style_map.get(style_choice, style_choice)
    
    if target_style not in available_styles:
        print(f"⚠️  Style '{target_style}' not available. Will try to match or use manual formatting.")
    
    # Find all tables
    print("🔍 Finding all tables...")
    all_tables = get_all_tables_recursive(doc)
    print(f"📊 Found {len(all_tables)} table(s)")
    
    # Process each table
    success_count = 0
    manual_count = 0
    
    for i, table in enumerate(all_tables, 1):
        print(f"  Processing table {i}/{len(all_tables)}...", end=" ")
        
        # Try to apply style first
        if target_style and target_style in available_styles:
            success, message = apply_best_style(table, target_style, available_styles)
            if success:
                success_count += 1
                print(f"✅ {message}")
                continue
        
        # If style fails, apply manual formatting
        success, message = apply_manual_formatting(table)
        if success:
            manual_count += 1
            print(f"✨ {message}")
        else:
            print(f"❌ {message}")
    
    # Save document
    doc.save(output_path)
    print(f"\n💾 Saved to: {output_path}")
    print(f"📊 Summary: {success_count} tables styled, {manual_count} manually formatted")
    
    return True

def main():
    print("=" * 60)
    print("📝 DOCX TABLE STYLER - USING AVAILABLE STYLES")
    print("=" * 60)
    print()
    
    # Get input file
    input_file = input("Enter DOCX path: ").strip().strip('"\'')
    
    if not os.path.exists(input_file):
        print("❌ File not found!")
        return
    
    # Preview available styles
    temp_doc = Document(input_file)
    available = [s.name for s in temp_doc.styles if s.type == WD_STYLE_TYPE.TABLE]
    
    print(f"\n📋 Available styles in your document:")
    for i, style in enumerate(available, 1):
        print(f"  {i}. '{style}'")
    
    print("\n🎨 Choose styling option:")
    print("  1. Use 'Normal Table' (professional look)")
    print("  2. Use 'Plain Table 3' (minimalist)")
    print("  3. Auto-select best available")
    print("  4. Manual formatting only")
    print("  5. Enter custom style name")
    
    choice = input("\nYour choice (1-5): ").strip()
    
    style_map = {
        "1": "normal",
        "2": "plain", 
        "3": "auto",
        "4": "manual"
    }
    
    if choice in style_map:
        style_choice = style_map[choice]
    elif choice == "5":
        style_choice = input("Enter exact style name: ").strip()
    else:
        style_choice = "auto"
        print("Using auto-select")
    
    # Output file
    base, ext = os.path.splitext(input_file)
    output_file = f"{base}_styled{ext}"
    
    print(f"\n🎯 Processing with option: {style_choice}")
    print(f"📂 Output will be: {output_file}")
    print()
    
    # Process the document
    success = process_document(input_file, output_file, style_choice)
    
    if success:
        print("\n✨ Document processing complete!")
    else:
        print("\n❌ Processing failed")

if __name__ == "__main__":
    main()