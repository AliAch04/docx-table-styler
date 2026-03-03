"""
Simple table formatter - just makes all tables look nice
"""

from docx import Document
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor

def find_tables(element, tables=None):
    """Find all tables recursively"""
    if tables is None:
        tables = []
    
    if hasattr(element, 'tables'):
        for table in element.tables:
            if table not in tables:
                tables.append(table)
                for row in table.rows:
                    for cell in row.cells:
                        find_tables(cell, tables)
    return tables

def format_table_nicely(table):
    """Apply nice consistent formatting to a table"""
    
    # Add borders to all cells
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Add borders
            borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                               r'<w:top w:val="single" w:sz="4"/>'
                               r'<w:left w:val="single" w:sz="4"/>'
                               r'<w:bottom w:val="single" w:sz="4"/>'
                               r'<w:right w:val="single" w:sz="4"/>'
                               r'</w:tcBorders>')
            tcPr.append(borders)
            
            # Set font size
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    # Format header row
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            # Add shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                               r'w:val="clear" w:color="auto" w:fill="D9E1F2"/>')
            tcPr.append(shading)
            
            # Make text bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def main():
    path = input("Enter DOCX path: ").strip('"\'')
    
    print("Loading document...")
    doc = Document(path)
    
    print("Finding tables...")
    tables = find_tables(doc)
    print(f"Found {len(tables)} tables")
    
    print("Formatting tables...")
    for i, table in enumerate(tables, 1):
        print(f"  Table {i}/{len(tables)}")
        try:
            format_table_nicely(table)
        except Exception as e:
            print(f"    Warning: {e}")
    
    output = path.replace('.docx', '_formatted.docx')
    doc.save(output)
    print(f"\n✅ Done! Saved to: {output}")

if __name__ == "__main__":
    main()