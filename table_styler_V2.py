"""
DOCX Table Style Changer - FIXED VERSION for nested tables
Handles complex documents with tables inside tables
"""

import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml

def get_all_tables_recursive(element, tables_list=None):
    """
    Recursively find ALL tables including nested ones
    
    Args:
        element: A document element (could be document, table, cell, etc.)
        tables_list: List to accumulate tables
    
    Returns:
        List of all tables found recursively
    """
    if tables_list is None:
        tables_list = []
    
    # If this element itself is a table, add it
    if hasattr(element, 'tables'):
        # For document objects
        for table in element.tables:
            if table not in tables_list:
                tables_list.append(table)
                # Recursively search inside this table's cells
                for row in table.rows:
                    for cell in row.cells:
                        get_all_tables_recursive(cell, tables_list)
    
    # If this is a cell, it might contain tables
    if hasattr(element, 'tables'):
        for table in element.tables:
            if table not in tables_list:
                tables_list.append(table)
                # Recursively search inside this nested table
                get_all_tables_recursive(table, tables_list)
    
    return tables_list

def list_document_styles(docx_path):
    """Show all styles available in the document"""
    try:
        doc = Document(docx_path)
        print(f"\n📋 Styles available in '{os.path.basename(docx_path)}':")
        print("-" * 50)
        
        # Get all table styles
        table_styles = [s for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE]
        
        if table_styles:
            print("TABLE STYLES:")
            for i, style in enumerate(table_styles, 1):
                print(f"  {i}. '{style.name}'")
        else:
            print("No table styles found in document!")
            
        # Also show built-in styles that might work
        print("\n💡 Built-in styles you can try:")
        built_in = [
            "Table Grid",
            "Light Shading Accent 1",
            "Light Shading Accent 2",
            "Light Shading Accent 3",
            "Medium Shading 1 Accent 1",
            "Medium Shading 1 Accent 2",
            "Medium Shading 2 Accent 1",
            "Medium Shading 2 Accent 2",
            "Light List Accent 1",
            "Light List Accent 2",
            "Light Grid Accent 1",
            "Light Grid Accent 2",
            "Medium Grid 1 Accent 1",
            "Medium Grid 1 Accent 2",
            "Medium Grid 2 Accent 1",
            "Medium Grid 2 Accent 2",
            "Dark List Accent 1",
            "Dark List Accent 2",
            "Colorful Shading Accent 1",
            "Colorful Shading Accent 2",
            "Colorful List Accent 1",
            "Colorful List Accent 2"
        ]
        for style in built_in:
            print(f"  • '{style}'")
            
        return table_styles
    except Exception as e:
        print(f"❌ Error reading document: {e}")
        return []

def apply_table_style_failsafe(table, style_name):
    """
    Apply style with fallback methods if direct assignment fails
    """
    try:
        # Method 1: Direct assignment (simplest)
        table.style = style_name
        return True, "Direct style applied"
    except:
        pass
    
    try:
        # Method 2: Try with different case variations
        variations = [
            style_name,
            style_name.lower(),
            style_name.upper(),
            style_name.title(),
            style_name.replace(' ', ''),
            'Table' + style_name.replace(' ', '')
        ]
        
        for var in variations:
            try:
                table.style = var
                return True, f"Style applied with variation: '{var}'"
            except:
                continue
    except:
        pass
    
    try:
        # Method 3: Apply basic formatting manually
        # This is a fallback that applies simple formatting
        for row in table.rows:
            for cell in row.cells:
                # Set borders
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add simple borders
                borders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                                   r'<w:top w:val="single" w:sz="4"/>'
                                   r'<w:left w:val="single" w:sz="4"/>'
                                   r'<w:bottom w:val="single" w:sz="4"/>'
                                   r'<w:right w:val="single" w:sz="4"/>'
                                   r'</w:tcBorders>')
                tcPr.append(borders)
        
        # Apply light shading to header row
        if len(table.rows) > 0:
            for cell in table.rows[0].cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                   r'w:val="clear" w:color="auto" w:fill="D9E1F2"/>')
                tcPr.append(shading)
        
        return True, "Manual formatting applied"
    except Exception as e:
        return False, f"All methods failed: {e}"

def fix_table_styles(input_path, output_path, style_name="Table Grid"):
    """
    Main function to fix table styles in complex documents
    """
    try:
        print(f"📂 Loading: {input_path}")
        doc = Document(input_path)
        
        # Get ALL tables recursively
        print("🔍 Searching for all tables (including nested)...")
        all_tables = get_all_tables_recursive(doc)
        
        print(f"📊 Found {len(all_tables)} table(s) total")
        
        # Statistics
        success_count = 0
        fail_count = 0
        
        # Apply style to each table
        for i, table in enumerate(all_tables, 1):
            print(f"  Processing table {i}/{len(all_tables)}...", end=" ")
            
            success, message = apply_table_style_failsafe(table, style_name)
            
            if success:
                success_count += 1
                print(f"✅ {message}")
            else:
                fail_count += 1
                print(f"❌ {message}")
        
        # Save the modified document
        doc.save(output_path)
        print(f"\n💾 Saved to: {output_path}")
        print(f"📊 Summary: {success_count} tables styled successfully, {fail_count} failed")
        
        return True
        
    except FileNotFoundError:
        print(f"❌ Error: File '{input_path}' not found")
        return False
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        return False

def main():
    """Interactive main function"""
    print("=" * 60)
    print("📝 DOCX TABLE STYLE CHANGER - FIXED VERSION")
    print("=" * 60)
    print("This version handles nested tables and complex documents")
    print()
    
    # Get input file
    input_file = input("Enter path to your DOCX file: ").strip().strip('"\'')
    
    if not os.path.exists(input_file):
        print("❌ File doesn't exist!")
        return
    
    # First, show available styles in this document
    list_document_styles(input_file)
    
    # Get style preference
    print("\n🎨 Enter the style name you want to apply")
    print("   (or press Enter for 'Table Grid')")
    style_name = input("Style: ").strip()
    
    if not style_name:
        style_name = "Table Grid"
        print(f"Using default: '{style_name}'")
    
    # Generate output filename
    base, ext = os.path.splitext(input_file)
    default_output = f"{base}_styled{ext}"
    
    output_file = input(f"\n📂 Output path (default: {default_output}): ").strip().strip('"\'')
    if not output_file:
        output_file = default_output
    
    print(f"\n🎯 Applying style: '{style_name}'")
    print("This may take a moment for complex documents...\n")
    
    # Apply the changes
    success = fix_table_styles(input_file, output_file, style_name)
    
    if success:
        print("\n✨ Done! Your document has been processed.")
        print(f"   Check: {output_file}")
    else:
        print("\n❌ Failed to process the document.")

if __name__ == "__main__":
    main()