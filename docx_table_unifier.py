"""
DOCX TABLE UNIFIER - One script to rule all table styling
Handles any DOCX file gracefully with user theme selection
"""

import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
import sys

class DocxTableUnifier:
    """Main class to handle DOCX table styling"""
    
    # Predefined themes with visual descriptions
    THEMES = {
        "1": {
            "name": "Professional Blue",
            "style": "Normal Table",
            "description": "Clean, professional with blue headers",
            "header_bg": "D9E1F2",  # Light blue
            "border_style": "single",
            "font_size": 10,
            "header_bold": True
        },
        "2": {
            "name": "Minimalist Light",
            "style": "Plain Table 3",
            "description": "Minimal, clean with light borders",
            "header_bg": "F2F2F2",  # Light gray
            "border_style": "single",
            "font_size": 10,
            "header_bold": True
        },
        "3": {
            "name": "Modern Grid",
            "style": "Table Grid",
            "description": "Clear grid with alternating rows (may need manual formatting)",
            "header_bg": "E6F0FA",  # Very light blue
            "border_style": "single",
            "font_size": 10,
            "header_bold": True,
            "alternating_rows": True
        },
        "4": {
            "name": "Academic Report",
            "style": "Light Shading Accent 1",
            "description": "Traditional academic table style",
            "header_bg": "4472C4",  # Dark blue
            "header_text": "FFFFFF",  # White text
            "border_style": "single",
            "font_size": 11,
            "header_bold": True
        },
        "5": {
            "name": "Corporate Dark",
            "style": "Medium Shading 1 Accent 2",
            "description": "Bold, corporate look",
            "header_bg": "44546A",  # Dark blue-gray
            "header_text": "FFFFFF",  # White text
            "border_style": "single",
            "font_size": 10,
            "header_bold": True
        },
        "6": {
            "name": "Simple Borders Only",
            "style": None,  # Manual formatting only
            "description": "Just add borders, keep existing formatting",
            "header_bg": None,
            "border_style": "single",
            "font_size": None
        },
        "7": {
            "name": "Custom Style",
            "style": "custom",
            "description": "Enter any style name manually",
            "header_bg": "D9E1F2",
            "border_style": "single",
            "font_size": 10
        }
    }
    
    def __init__(self, docx_path):
        """Initialize with document path"""
        self.docx_path = docx_path
        self.doc = None
        self.available_styles = []
        self.tables = []
        
    def load_document(self):
        """Load the document and analyze its contents"""
        try:
            print(f"📂 Loading: {os.path.basename(self.docx_path)}")
            self.doc = Document(self.docx_path)
            
            # Get available table styles
            self.available_styles = [s.name for s in self.doc.styles 
                                    if s.type == WD_STYLE_TYPE.TABLE]
            
            # Find all tables (including nested)
            self.tables = self._find_all_tables(self.doc)
            
            print(f"✅ Found {len(self.tables)} table(s)")
            print(f"📋 Available styles: {len(self.available_styles)} found")
            return True
            
        except Exception as e:
            print(f"❌ Error loading document: {e}")
            return False
    
    def _find_all_tables(self, element, tables=None):
        """Recursively find all tables in the document"""
        if tables is None:
            tables = []
        
        if hasattr(element, 'tables'):
            for table in element.tables:
                if table not in tables:
                    tables.append(table)
                    # Search inside table cells for nested tables
                    for row in table.rows:
                        for cell in row.cells:
                            self._find_all_tables(cell, tables)
        return tables
    
    def display_available_styles(self):
        """Show what styles are actually available in this document"""
        print("\n📋 STYLES AVAILABLE IN THIS DOCUMENT:")
        print("-" * 40)
        if self.available_styles:
            for i, style in enumerate(self.available_styles, 1):
                print(f"  {i}. '{style}'")
        else:
            print("  ⚠️  No predefined table styles found")
            print("  Will use manual formatting only")
        print()
    
    def display_theme_menu(self):
        """Show available themes to user"""
        print("\n🎨 AVAILABLE THEMES:")
        print("=" * 60)
        for key, theme in self.THEMES.items():
            print(f"\n  {key}. {theme['name']}")
            print(f"     📝 {theme['description']}")
            if key == "1" and self.available_styles:
                if theme['style'] in self.available_styles:
                    print(f"     ✅ Available in your document")
                else:
                    print(f"     ⚠️  Will use manual formatting")
        print("\n  Enter 'list' to see all available styles in your document")
        print("  Enter 'q' to quit")
        print("=" * 60)
    
    def apply_theme_to_table(self, table, theme):
        """Apply selected theme to a single table with fallbacks"""
        results = []
        
        # Method 1: Try to apply style if specified and available
        if theme.get('style') and theme['style'] != 'custom':
            if theme['style'] in self.available_styles:
                try:
                    table.style = theme['style']
                    results.append(f"✅ Applied style: '{theme['style']}'")
                except:
                    results.append(f"⚠️  Could not apply style '{theme['style']}'")
            else:
                results.append(f"⚠️  Style '{theme['style']}' not available")
        
        # Method 2: Apply manual formatting based on theme
        try:
            self._apply_manual_formatting(table, theme)
            results.append("✨ Manual formatting applied")
        except Exception as e:
            results.append(f"❌ Manual formatting failed: {e}")
        
        return " | ".join(results)
    
    def _apply_manual_formatting(self, table, theme):
        """Apply manual formatting based on theme settings"""
        
        # Set table properties
        table.autofit = True
        table.allow_autofit = True
        
        # Process each cell
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add borders
                if theme.get('border_style'):
                    borders_xml = f'''<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:top w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                        <w:left w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                        <w:bottom w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                        <w:right w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                    </w:tcBorders>'''
                    borders = parse_xml(borders_xml)
                    tcPr.append(borders)
                
                # Center align vertically
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Format text
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    
                    for run in paragraph.runs:
                        if theme.get('font_size'):
                            run.font.size = Pt(theme['font_size'])
                
                # Header row special formatting
                if row_idx == 0 and theme.get('header_bg'):
                    # Add background color to header
                    bg_color = theme['header_bg']
                    if bg_color:
                        shading_xml = f'''<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                            w:val="clear" w:color="auto" w:fill="{bg_color}"/>'''
                        shading = parse_xml(shading_xml)
                        tcPr.append(shading)
                    
                    # Make header text bold
                    if theme.get('header_bold', True):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Header text color
                    if theme.get('header_text'):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor.from_string(theme['header_text'])
        
        # Alternating row colors if specified
        if theme.get('alternating_rows') and len(table.rows) > 1:
            for row_idx, row in enumerate(table.rows):
                if row_idx % 2 == 1 and row_idx > 0:  # Alternate rows (skip header)
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                           r'w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
                        tcPr.append(shading)
    
    def process_document(self, theme_choice, output_path=None):
        """Process the entire document with chosen theme"""
        
        # Get theme configuration
        if theme_choice == "7":  # Custom style
            custom_style = input("\n📝 Enter exact style name: ").strip()
            theme = self.THEMES["7"].copy()
            theme['style'] = custom_style
        elif theme_choice in self.THEMES:
            theme = self.THEMES[theme_choice].copy()
        else:
            print("❌ Invalid theme choice")
            return False
        
        # Generate output path if not provided
        if not output_path:
            base, ext = os.path.splitext(self.docx_path)
            output_path = f"{base}_unified{ext}"
        
        print(f"\n🎨 Applying theme: {theme['name']}")
        print(f"📝 {theme['description']}")
        print("-" * 50)
        
        # Apply theme to each table
        success_count = 0
        for i, table in enumerate(self.tables, 1):
            print(f"  Table {i}/{len(self.tables)}...", end=" ")
            result = self.apply_theme_to_table(table, theme)
            print(result)
            if "✅" in result or "✨" in result:
                success_count += 1
        
        # Save document
        self.doc.save(output_path)
        print(f"\n💾 Saved to: {output_path}")
        print(f"📊 Summary: {success_count}/{len(self.tables)} tables processed successfully")
        
        return True

def main():
    """Main interactive function"""
    print("=" * 70)
    print("📝 DOCX TABLE UNIFIER - Professional Table Styling Tool")
    print("=" * 70)
    print("This tool will unify all tables in your document with a consistent theme")
    print()
    
    # Get input file
    while True:
        input_file = input("📂 Enter path to DOCX file: ").strip().strip('"\'')
        if os.path.exists(input_file):
            break
        print("❌ File not found! Please try again.")
    
    # Initialize unifier
    unifier = DocxTableUnifier(input_file)
    
    # Load document
    if not unifier.load_document():
        print("❌ Failed to load document. Exiting.")
        return
    
    # Main interaction loop
    while True:
        # Show available styles in this document
        unifier.display_available_styles()
        
        # Show theme menu
        unifier.display_theme_menu()
        
        # Get user choice
        choice = input("\n📝 Your choice: ").strip().lower()
        
        if choice == 'q':
            print("👋 Goodbye!")
            return
        
        if choice == 'list':
            unifier.display_available_styles()
            input("\nPress Enter to continue...")
            continue
        
        if choice in unifier.THEMES:
            # Generate output filename
            base, ext = os.path.splitext(input_file)
            output_file = f"{base}_unified{ext}"
            
            # Process document
            if unifier.process_document(choice, output_file):
                print("\n✅ Document processed successfully!")
                
                # Ask if user wants to process another file
                again = input("\n🔄 Process another file? (y/n): ").strip().lower()
                if again == 'y':
                    # Start over with new file
                    main()
                else:
                    print("\n👋 Thank you for using DOCX Table Unifier!")
                return
            else:
                print("\n❌ Failed to process document.")
        else:
            print("\n❌ Invalid choice. Please try again.")
            input("Press Enter to continue...")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n👋 Goodbye!")
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
        print("Please report this issue with your document.")