"""
DOCX Table Style Changer
A tool to change the style of all tables in a Word document.
"""

import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

def list_available_styles():
    """Display all available table styles in Word"""
    doc = Document()
    table_styles = [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE]
    print("\n📋 Available Table Styles:")
    print("-" * 40)
    for i, style in enumerate(table_styles, 1):
        print(f"{i}. {style}")
    return table_styles

def change_table_style(input_path, output_path, style_name):
    """
    Change all tables in a document to the specified style
    
    Args:
        input_path (str): Path to input DOCX file
        output_path (str): Path to save modified DOCX file
        style_name (str): Name of the table style to apply
    """
    try:
        # Load the document
        print(f"📂 Loading: {input_path}")
        doc = Document(input_path)
        
        # Get all tables
        tables = doc.tables
        print(f"📊 Found {len(tables)} table(s)")
        
        # Apply new style to each table
        for i, table in enumerate(tables, 1):
            try:
                table.style = style_name
                print(f"  ✅ Table {i}: Style applied")
            except Exception as e:
                print(f"  ❌ Table {i}: Failed - {e}")
        
        # Save the modified document
        doc.save(output_path)
        print(f"💾 Saved to: {output_path}")
        return True
        
    except FileNotFoundError:
        print(f"❌ Error: File '{input_path}' not found")
        return False
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        return False

def main():
    """Main function to run the script interactively"""
    print("=" * 50)
    print("📝 DOCX TABLE STYLE CHANGER")
    print("=" * 50)
    
    # Show available styles first
    styles = list_available_styles()
    
    # Get user input
    print("\n📂 Enter your file paths:")
    input_file = input("Input DOCX path: ").strip()
    
    # Remove quotes if user pasted with quotes
    input_file = input_file.strip('"\'')
    
    if not os.path.exists(input_file):
        print("❌ File doesn't exist!")
        return
    
    # Generate default output filename
    base, ext = os.path.splitext(input_file)
    default_output = f"{base}_styled{ext}"
    
    output_file = input(f"Output path (default: {default_output}): ").strip().strip('"\'')
    if not output_file:
        output_file = default_output
    
    # Get style
    print(f"\n🎨 Choose style (1-{len(styles)} or type the name):")
    style_input = input("Style: ").strip()
    
    # Check if user entered a number
    try:
        style_index = int(style_input) - 1
        if 0 <= style_index < len(styles):
            style_name = styles[style_index]
        else:
            print("❌ Invalid number, using first style")
            style_name = styles[0]
    except ValueError:
        # User entered a style name
        style_name = style_input
    
    print(f"\n🎯 Applying style: '{style_name}'")
    
    # Apply the changes
    success = change_table_style(input_file, output_file, style_name)
    
    if success:
        print("\n✨ Done! Your document is ready.")
    else:
        print("\n❌ Failed to process the document.")

if __name__ == "__main__":
    main()