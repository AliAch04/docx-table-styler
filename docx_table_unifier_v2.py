"""
DOCX TABLE UNIFIER V2 - Smart Theme Mapping
Now intelligently uses available styles from your document!
"""

import os
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL

class DocxTableUnifierV2:
    """Smart table unifier that adapts to available styles"""
    
    # Theme definitions with flexible style mapping
    THEMES = {
        "1": {
            "name": "Professional Blue",
            "description": "Clean, professional with blue headers",
            "preferred_styles": ["Normal Table", "Light Shading Accent 1", "Medium Shading 1 Accent 1"],
            "fallback_style": None,  # Will use first available
            "header_bg": "D9E1F2",
            "border_style": "single",
            "font_size": 10,
            "header_bold": True
        },
        "2": {
            "name": "Minimalist Light",
            "description": "Minimal, clean with light borders",
            "preferred_styles": ["Plain Table 3", "Light List Accent 1", "Light Grid Accent 1"],
            "fallback_style": None,
            "header_bg": "F2F2F2",
            "border_style": "single",
            "font_size": 10,
            "header_bold": True
        },
        "3": {
            "name": "Modern Grid",
            "description": "Clear grid with alternating rows",
            "preferred_styles": ["Table Grid", "Medium Grid 1 Accent 1", "Light Grid Accent 1"],
            "fallback_style": None,
            "header_bg": "E6F0FA",
            "border_style": "single",
            "font_size": 10,
            "header_bold": True,
            "alternating_rows": True
        },
        "4": {
            "name": "Academic Report",
            "description": "Traditional academic table style",
            "preferred_styles": ["Light Shading Accent 1", "Medium Shading 1 Accent 1", "Normal Table"],
            "fallback_style": None,
            "header_bg": "4472C4",
            "header_text": "FFFFFF",
            "border_style": "single",
            "font_size": 11,
            "header_bold": True
        },
        "5": {
            "name": "Corporate Dark",
            "description": "Bold, corporate look",
            "preferred_styles": ["Medium Shading 1 Accent 2", "Dark List Accent 1", "Normal Table"],
            "fallback_style": None,
            "header_bg": "44546A",
            "header_text": "FFFFFF",
            "border_style": "single",
            "font_size": 10,
            "header_bold": True
        },
        "6": {
            "name": "Simple Borders Only",
            "description": "Just add borders, keep existing formatting",
            "preferred_styles": [],
            "fallback_style": None,
            "header_bg": None,
            "border_style": "single",
            "font_size": None
        },
        "7": {
            "name": "Custom Style",
            "description": "Enter any style name manually",
            "preferred_styles": [],
            "fallback_style": "custom",
            "header_bg": "D9E1F2",
            "border_style": "single",
            "font_size": 10
        }
    }
    
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = None
        self.available_styles = []
        self.tables = []
        self.style_mapping = {}  # Maps theme names to actual styles
        
    def load_document(self):
        """Load and analyze document"""
        try:
            print(f"📂 Loading: {os.path.basename(self.docx_path)}")
            self.doc = Document(self.docx_path)
            
            # Get ALL available table styles
            self.available_styles = [s.name for s in self.doc.styles 
                                    if s.type == WD_STYLE_TYPE.TABLE]
            
            # Find all tables
            self.tables = self._find_all_tables(self.doc)
            
            print(f"✅ Found {len(self.tables)} table(s)")
            print(f"📋 Available styles: {len(self.available_styles)} found")
            
            # Show them immediately
            self._display_available_styles()
            
            return True
        except Exception as e:
            print(f"❌ Error: {e}")
            return False
    
    def _find_all_tables(self, element, tables=None):
        """Recursively find tables"""
        if tables is None:
            tables = []
        
        if hasattr(element, 'tables'):
            for table in element.tables:
                if table not in tables:
                    tables.append(table)
                    for row in table.rows:
                        for cell in row.cells:
                            self._find_all_tables(cell, tables)
        return tables
    
    def _display_available_styles(self):
        """Show available styles with descriptions"""
        print("\n" + "="*60)
        print("📋 STYLES AVAILABLE IN YOUR DOCUMENT")
        print("="*60)
        
        if self.available_styles:
            # Group styles by type
            print("\n🎨 Table Styles Found:")
            for i, style in enumerate(self.available_styles, 1):
                # Add visual indicators
                if "Normal" in style:
                    indicator = "📊 (Basic)"
                elif "Plain" in style:
                    indicator = "✨ (Minimal)"
                elif "Light" in style:
                    indicator = "🌟 (Light)"
                elif "Medium" in style:
                    indicator = "⭐ (Medium)"
                elif "Dark" in style:
                    indicator = "🌙 (Dark)"
                elif "Grid" in style:
                    indicator = "🔲 (Grid)"
                elif "List" in style:
                    indicator = "📋 (List)"
                else:
                    indicator = "📌"
                
                print(f"  {i:2d}. '{style}' {indicator}")
            
            print("\n💡 TIP: These are the ONLY styles that will work directly!")
            print("   Other style names will trigger manual formatting.")
        else:
            print("  ⚠️  No predefined table styles found")
            print("  Will use manual formatting only")
    
    def _find_best_style_match(self, theme_preferences):
        """Find the best matching available style"""
        if not theme_preferences:
            return None
            
        # Try each preferred style in order
        for preferred in theme_preferences:
            for available in self.available_styles:
                if preferred.lower() in available.lower() or available.lower() in preferred.lower():
                    print(f"  🔍 Matched: '{preferred}' → '{available}'")
                    return available
        
        # If no match, try first available
        if self.available_styles:
            print(f"  🔍 No exact match, using first available: '{self.available_styles[0]}'")
            return self.available_styles[0]
        
        return None
    
    def display_theme_menu(self):
        """Show themes with actual style mapping"""
        print("\n" + "="*70)
        print("🎨 AVAILABLE THEMES (Smart-Mapped to Your Document)")
        print("="*70)
        
        for key, theme in self.THEMES.items():
            print(f"\n  {key}. {theme['name']}")
            print(f"     📝 {theme['description']}")
            
            # Find best style for this theme
            best_style = self._find_best_style_match(theme['preferred_styles'])
            
            if best_style:
                print(f"     ✅ Will use: '{best_style}'")
                # Store the mapping
                self.style_mapping[key] = best_style
            else:
                if theme['name'] == "Custom Style":
                    print(f"     ✏️  You'll enter custom style name")
                elif theme['name'] == "Simple Borders Only":
                    print(f"     🔲 Manual formatting only (no style)")
                else:
                    print(f"     ⚠️  No matching style - will use manual formatting")
                    self.style_mapping[key] = None
        
        print("\n" + "-"*70)
        print("  'list' - Show all available styles again")
        print("  'q'    - Quit")
        print("-"*70)
    
    def apply_theme_to_table(self, table, theme_key, custom_style=None):
        """Apply theme with smart style selection"""
        
        theme = self.THEMES[theme_key]
        results = []
        
        # Get the mapped style
        if theme_key == "7" and custom_style:  # Custom style
            style_to_use = custom_style
        else:
            style_to_use = self.style_mapping.get(theme_key)
        
        # Try to apply style if available
        if style_to_use and style_to_use in self.available_styles:
            try:
                table.style = style_to_use
                results.append(f"✅ Applied style: '{style_to_use}'")
            except Exception as e:
                results.append(f"⚠️  Style application failed: {e}")
        else:
            if style_to_use:
                results.append(f"⚠️  Style '{style_to_use}' not available")
        
        # Always apply manual formatting enhancements
        try:
            self._apply_manual_formatting(table, theme)
            results.append("✨ Enhanced with theme formatting")
        except Exception as e:
            results.append(f"❌ Formatting failed: {e}")
        
        return " | ".join(results)
    
    def _apply_manual_formatting(self, table, theme):
        """Apply theme-specific manual formatting"""
        
        # Basic table setup
        table.autofit = True
        table.allow_autofit = True
        
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add borders
                if theme.get('border_style'):
                    borders_xml = f'''<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:top w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                        <w:left w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                        <w:bottom w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                        <w:right w:val="{theme['border_style']}" w:sz="4" w:space="0" w:color="auto"/>
                    </w:tcBorders>'''
                    borders = parse_xml(borders_xml)
                    tcPr.append(borders)
                
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # Text formatting
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    
                    for run in paragraph.runs:
                        if theme.get('font_size'):
                            run.font.size = Pt(theme['font_size'])
                
                # Header row
                if row_idx == 0 and theme.get('header_bg'):
                    # Background
                    if theme['header_bg']:
                        shading_xml = f'''<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                            w:val="clear" w:color="auto" w:fill="{theme['header_bg']}"/>'''
                        shading = parse_xml(shading_xml)
                        tcPr.append(shading)
                    
                    # Bold
                    if theme.get('header_bold', True):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    
                    # Text color
                    if theme.get('header_text'):
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.color.rgb = RGBColor.from_string(theme['header_text'])
        
        # Alternating rows
        if theme.get('alternating_rows') and len(table.rows) > 1:
            for row_idx, row in enumerate(table.rows):
                if row_idx % 2 == 1 and row_idx > 0:
                    for cell in row.cells:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shading = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                                           r'w:val="clear" w:color="auto" w:fill="F5F5F5"/>')
                        tcPr.append(shading)
    
    def process_document(self, theme_key, output_path=None, custom_style=None):
        """Process document with smart theme"""
        
        theme = self.THEMES[theme_key]
        
        if not output_path:
            base, ext = os.path.splitext(self.docx_path)
            output_path = f"{base}_unified{ext}"
        
        print(f"\n🎨 Applying theme: {theme['name']}")
        
        # Show what style we're using
        if theme_key == "7" and custom_style:
            print(f"📝 Custom style: '{custom_style}'")
        elif theme_key in self.style_mapping and self.style_mapping[theme_key]:
            print(f"📝 Using style: '{self.style_mapping[theme_key]}'")
        else:
            print(f"📝 Using manual formatting with theme colors")
        
        print("-" * 50)
        
        # Process each table
        success_count = 0
        for i, table in enumerate(self.tables, 1):
            print(f"  Table {i}/{len(self.tables)}...", end=" ")
            result = self.apply_theme_to_table(table, theme_key, custom_style)
            print(result)
            if "✅" in result or "✨" in result:
                success_count += 1
        
        # Save
        self.doc.save(output_path)
        print(f"\n💾 Saved to: {output_path}")
        print(f"📊 Summary: {success_count}/{len(self.tables)} tables processed")
        
        return True

def main():
    """Main function"""
    print("=" * 70)
    print("📝 DOCX TABLE UNIFIER V2 - Smart Theme Mapping")
    print("=" * 70)
    print("This tool intelligently maps themes to styles available in YOUR document")
    print()
    
    # Get file
    while True:
        input_file = input("📂 Enter path to DOCX file: ").strip().strip('"\'')
        if os.path.exists(input_file):
            break
        print("❌ File not found!")
    
    # Initialize
    unifier = DocxTableUnifierV2(input_file)
    
    if not unifier.load_document():
        return
    
    # Interactive loop
    while True:
        unifier.display_theme_menu()
        
        choice = input("\n📝 Your choice: ").strip().lower()
        
        if choice == 'q':
            print("👋 Goodbye!")
            return
        
        if choice == 'list':
            unifier._display_available_styles()
            input("\nPress Enter to continue...")
            continue
        
        if choice in unifier.THEMES:
            custom_style = None
            if choice == "7":  # Custom style
                print("\n📝 Available styles in your document:")
                for i, style in enumerate(unifier.available_styles, 1):
                    print(f"  {i}. '{style}'")
                custom_style = input("\nEnter exact style name: ").strip()
            
            # Process
            unifier.process_document(choice, custom_style=custom_style)
            
            # Ask for another
            again = input("\n🔄 Process another file? (y/n): ").strip().lower()
            if again == 'y':
                main()
            else:
                print("\n👋 Thank you for using DOCX Table Unifier V2!")
            return
        else:
            print("❌ Invalid choice")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n👋 Goodbye!")
    except Exception as e:
        print(f"\n❌ Error: {e}")