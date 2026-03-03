from docx import Document

def debug_document_structure(docx_path):
    """Analyze document structure"""
    print(f"\n🔍 Analyzing: {docx_path}")
    
    doc = Document(docx_path)
    
    # Method 1: Direct tables
    direct_tables = len(doc.tables)
    print(f"📊 Direct tables (doc.tables): {direct_tables}")
    
    # Method 2: Try to find all tables recursively
    def count_tables_recursive(element, level=0):
        count = 0
        if hasattr(element, 'tables'):
            for table in element.tables:
                count += 1
                print(f"  {'  ' * level}Found table at level {level}")
                for row in table.rows:
                    for cell in row.cells:
                        count += count_tables_recursive(cell, level + 1)
        return count
    
    total_tables = count_tables_recursive(doc)
    print(f"📊 Total tables (including nested): {total_tables}")
    
    # Check available styles
    from docx.enum.style import WD_STYLE_TYPE
    table_styles = [s.name for s in doc.styles if s.type == WD_STYLE_TYPE.TABLE]
    print(f"\n🎨 Table styles in document: {len(table_styles)}")
    for style in table_styles[:10]:  # Show first 10
        print(f"  • '{style}'")

if __name__ == "__main__":
    path = input("Enter DOCX path: ").strip('"\'')
    debug_document_structure(path)